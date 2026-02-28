#!/usr/bin/env python3
"""
TRPG 로그 → EPUB/DOCX 변환기
코코포리아 / Roll20 지원
"""

from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os
import uuid
import yaml
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Callable
import mimetypes
import io
import logging

logger = logging.getLogger(__name__)

# ============ 이미지 최적화 ============
def optimize_image(img_path: Path, config: Dict[str, Any]) -> Tuple[bytes, str, str]:
    """이미지를 최적화하여 (bytes, mime_type, filename) 반환.

    - max_resolution 이상이면 리사이즈
    - JPEG 품질 조절
    - WebP → JPEG 변환 (옵션)
    """
    images_config = config.get('images', {})
    max_res = images_config.get('max_resolution', 0)
    quality = images_config.get('jpeg_quality', 85)
    convert_webp = images_config.get('convert_webp', True)

    try:
        from PIL import Image
    except ImportError:
        # Pillow 없으면 원본 그대로 반환
        with open(img_path, 'rb') as f:
            data = f.read()
        mime, _ = mimetypes.guess_type(str(img_path))
        return data, mime or 'image/jpeg', img_path.name

    img = Image.open(img_path)
    original_format = img.format or 'PNG'
    filename = img_path.name
    changed = False

    # WebP → JPEG 변환
    is_webp = original_format.upper() == 'WEBP' or str(img_path).lower().endswith('.webp')
    if is_webp and convert_webp:
        if img.mode in ('RGBA', 'LA', 'PA'):
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        original_format = 'JPEG'
        filename = img_path.stem + '.jpg'
        changed = True

    # 리사이즈
    if max_res and max_res > 0:
        w, h = img.size
        longest = max(w, h)
        if longest > max_res:
            ratio = max_res / longest
            new_w, new_h = int(w * ratio), int(h * ratio)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            changed = True

    # 저장
    buf = io.BytesIO()
    save_format = original_format.upper()
    if save_format in ('JPEG', 'JPG'):
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        mime_type = 'image/jpeg'
        if not filename.lower().endswith(('.jpg', '.jpeg')):
            filename = img_path.stem + '.jpg'
    elif save_format == 'PNG':
        img.save(buf, format='PNG', optimize=True)
        mime_type = 'image/png'
    else:
        # 기타 포맷 → JPEG로 변환
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        mime_type = 'image/jpeg'
        filename = img_path.stem + '.jpg'

    return buf.getvalue(), mime_type, filename


# ============ 기본 설정 ============
# ConfigManager.DEFAULT_CONFIG가 단일 소스 (Single Source of Truth)
# engine 단독 실행 시에만 사용되는 폴백
def _get_default_config():
    """기본 설정 반환 - ConfigManager가 있으면 그쪽을 사용"""
    try:
        from core.config_manager import ConfigManager
        return ConfigManager.DEFAULT_CONFIG.copy()
    except ImportError:
        pass
    # ConfigManager 임포트 실패 시 최소 폴백
    return {
        'paths': {'input_dir': './input', 'output_dir': './export', 'fonts_dir': './fonts', 'images_dir': './images'},
        'output_format': 'both', 'log_source': 'auto',
        'cover': {'image': '', 'include': True, 'title_on_cover': True, 'author_on_cover': True,
                  'background_color': '#1a1a1a', 'title_color': '#ffffff', 'subtitle': ''},
        'toc': {'include': True, 'title': '목차', 'mode': 'auto', 'entries': [], 'style': 'simple'},
        'fonts': {'name_font': "'Pretendard', sans-serif", 'body_font': "'Nanum Myeongjo', serif",
                  'pretendard_cdn': "https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css",
                  'embed': {}, 'docx_fallback': {'body': '맑은 고딕', 'name': '맑은 고딕'}},
        'style': {'narration_prefix': '＿', 'scene_marker': '■', 'dialogue_margin': 0.12,
                  'narration_margin': 0.8, 'narration_indent': 1.5, 'dice_color': '#888'},
        'narration': {'users': ['GM', 'KP', 'DM', 'Keeper', 'Narrator'], 'style': 'indent'},
        'content': {'include_dice': True, 'include_system': True, 'include_effects': True},
        'dialogue': {'merge_consecutive': False, 'merge_separator': '\n', 'merge_max': 5},
        'images': {'enable': True, 'markers': [r'\[IMG:\s*(.+?)\]', r'\[삽화:\s*(.+?)\]'], 'show_caption': True},
        'custom_styles': {},
        'chapter': {'split_mode': 'scene', 'entries_per_chapter': 300, 'extract_scene_title': True,
                    'title_format': '장면 {n}', 'min_scene_entries': 10, 'scene_patterns': ['^■', '^씬\\s*\\d+', '^장면\\s*\\d+']},
        'parsing': {'name_max_length': 50, 'skip_channels': [], 'normalize_punctuation': True},
    }

DEFAULT_CONFIG = _get_default_config()

def deep_merge(base: Dict, override: Dict) -> Dict:
    result = base.copy()
    for key, value in override.items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = value
    return result

def load_config(config_path: Optional[str] = None) -> Dict[str, Any]:
    config = DEFAULT_CONFIG.copy()
    if config_path is None:
        script_dir = Path(__file__).parent
        config_path = script_dir / 'config.yaml'
    if Path(config_path).exists():
        with open(config_path, 'r', encoding='utf-8') as f:
            user_config = yaml.safe_load(f)
            if user_config:
                config = deep_merge(config, user_config)
        logger.info("Config loaded: %s", config_path)
    return config

# ============ 유틸리티 ============
def get_font_files(config: Dict[str, Any]) -> Dict[str, Any]:
    fonts_dir = Path(config.get('paths', {}).get('fonts_dir', './fonts'))
    fonts = {'body': None, 'name': None, 'all': []}
    if not fonts_dir.exists():
        return fonts
    embed_config = config.get('fonts', {}).get('embed', {})
    for font_type in ['body', 'name']:
        filename = embed_config.get(font_type)
        if filename:
            font_path = fonts_dir / filename
            if font_path.exists():
                fonts[font_type] = font_path
                fonts['all'].append(font_path)
    return fonts

def get_font_family_name(font_path: Path) -> str:
    name = font_path.stem
    for suffix in ['-Regular', '-Bold', '-Light', '-Medium', '-SemiBold']:
        name = name.replace(suffix, '')
    return name

def normalize_punctuation(text: str) -> str:
    text = re.sub(r'\.{3,}', '……', text)
    text = text.replace('...', '……')
    return text

def escape_html(text: str) -> str:
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 3:
        hex_color = ''.join([c*2 for c in hex_color])
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def is_narration_user(name: str, config: Dict[str, Any]) -> bool:
    default_users = ['GM', 'KP', 'DM', 'System', '시스템', '나레이션', 'Narrator', '진행자', 'Storyteller', 'Keeper']
    config_users = config.get('narration', {}).get('users', [])
    # 기본값과 설정을 병합
    narration_users = list(set(default_users + config_users))
    return name.lower().strip() in [u.lower() for u in narration_users]

def match_custom_style(content: str, config: Dict[str, Any]) -> Optional[str]:
    for style_name, style_config in config.get('custom_styles', {}).items():
        pattern = style_config.get('pattern', '')
        if pattern:
            try:
                if re.search(pattern, content, re.IGNORECASE):
                    return style_config.get('type', 'dialogue')
            except re.error as e:
                logger.warning("커스텀 스타일 패턴 오류 '%s': %s", pattern, e)
    return None

def is_dice_roll(text: str) -> bool:
    """주사위 굴림 판정 - 다양한 형식 지원"""
    text_upper = text.upper()

    # 패턴 1: 코코포리아 스타일 (CCB<=65, 1D100<=50)
    if re.search(r'(CCB|1D100|2D6|\d+D\d+)', text_upper) and ('→' in text or '=' in text or '>' in text or '＞' in text):
        return True

    # 패턴 2: /roll 커맨드
    if re.search(r'/roll\s+\d*d\d+', text, re.IGNORECASE):
        return True

    # 패턴 3: Rolling XdY 형식
    if re.search(r'rolling\s+\d*d\d+', text, re.IGNORECASE):
        return True

    # 패턴 4: vs AC, vs DC 형식
    if re.search(r'\d+d\d+.*vs\s*(ac|dc)', text, re.IGNORECASE):
        return True

    # 패턴 5: Attack/Damage roll
    if re.search(r'(attack|damage|hit|save|check)\s*:?\s*\d*d\d+', text, re.IGNORECASE):
        return True

    # 패턴 6: 판정 결과 (성공/실패/크리티컬)
    if re.search(r'\d+d\d+', text, re.IGNORECASE) and re.search(r'(성공|실패|critical|crit|fumble|대성공|대실패)', text, re.IGNORECASE):
        return True

    # 패턴 7: 단순 주사위 결과 (3d6 = 12)
    if re.search(r'\d+d\d+\s*[+\-]?\s*\d*\s*=\s*\d+', text, re.IGNORECASE):
        return True

    return False

def validate_regex(pattern: str) -> bool:
    """정규식 패턴 유효성 검사"""
    try:
        re.compile(pattern)
        return True
    except re.error as e:
        logger.warning("잘못된 정규식 패턴 '%s': %s", pattern, e)
        return False

def is_scene_marker(text: str, patterns: List[str]) -> bool:
    for pattern in patterns:
        try:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        except re.error as e:
            logger.warning("씬 마커 패턴 오류 '%s': %s (리터럴 매칭 시도)", pattern, e)
            if pattern in text:
                return True
    return False

def extract_scene_title(text: str) -> Optional[str]:
    text = text.strip()
    text = re.sub(r'^[─━\-=\*]+\s*', '', text)
    text = re.sub(r'\s*[─━\-=\*]+$', '', text)
    return text.strip() if text else None

def find_image_file(filename: str, config: Dict[str, Any]) -> Optional[Path]:
    images_dir = Path(config.get('paths', {}).get('images_dir', './images'))
    if Path(filename).exists():
        return Path(filename)
    if images_dir.exists():
        img_path = images_dir / filename
        if img_path.exists():
            return img_path
    return None

def extract_image_markers(text: str, config: Dict[str, Any]) -> Tuple[Optional[str], str]:
    images_config = config.get('images', {})
    if not images_config.get('enable', True):
        return None, text
    markers = images_config.get('markers', [])
    for pattern in markers:
        try:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                filename = match.group(1).strip()
                remaining_text = re.sub(pattern, '', text, count=1).strip()
                return filename, remaining_text
        except re.error as e:
            logger.warning("이미지 마커 패턴 오류 '%s': %s", pattern, e)
    return None, text

# ============ 파싱 ============
def strip_channel_prefix(text: str) -> Tuple[str, Optional[str]]:
    """채널 접두사 제거 (예: [메인], [잡담], [정보] 등)"""
    # 코코포리아 스타일: [채널명] 텍스트
    match = re.match(r'^\[([^\]]+)\]\s*', text)
    if match:
        return text[match.end():].strip(), match.group(1)
    return text, None

def smart_split_name_content(text: str, max_name_length: int = 50) -> Tuple[Optional[str], str]:
    """이름과 내용을 지능적으로 분리 (복잡한 이름 처리)"""
    if ':' not in text:
        return None, text

    # 콜론 위치들 찾기
    colon_positions = [i for i, c in enumerate(text) if c == ':']

    if len(colon_positions) == 1:
        # 콜론이 하나면 시간 형식인지 확인 후 분리
        pos = colon_positions[0]
        if 0 < pos <= max_name_length:
            name_part = text[:pos].strip()
            # 시간 형식 (10:30, 9:15 등)이면 분리하지 않음
            if re.match(r'^\d{1,2}$', name_part):
                return None, text
            return name_part, text[pos+1:].strip()
        return None, text

    # 콜론이 여러 개인 경우 최적 분리점 찾기
    best_pos = None
    best_score = -1

    for pos in colon_positions:
        if pos == 0 or pos > max_name_length:
            continue

        name_part = text[:pos].strip()
        content_part = text[pos+1:].strip()

        # 점수 계산
        score = 0

        # 이름 부분이 적절한 길이면 가산점
        if 1 <= len(name_part) <= 30:
            score += 10

        # 이름에 괄호가 있으면 가산점 (캐릭터명 (직업) 형식)
        if '(' in name_part and ')' in name_part:
            score += 5

        # 콜론 뒤에 공백이 있으면 가산점
        if pos + 1 < len(text) and text[pos + 1] == ' ':
            score += 3

        # 콜론 앞에 공백이 있으면 가산점 (코코포리아 스타일: "이름 : 대사")
        if pos > 0 and text[pos - 1] == ' ':
            score += 2

        # 내용이 대화체로 시작하면 가산점
        if content_part and content_part[0] in '「『"\'""「':
            score += 5

        # 이름에 숫자나 특수기호만 있으면 감점
        if re.match(r'^[\d\W]+$', name_part):
            score -= 10

        # 이름에 시간 형식이 있으면 감점 (10:30 같은)
        if re.match(r'^\d{1,2}:\d{2}', name_part) or re.match(r'^\d{1,2}:\d{2}', text[:pos]):
            score -= 20

        if score > best_score:
            best_score = score
            best_pos = pos

    if best_pos is not None and best_score >= 0:
        return text[:best_pos].strip(), text[best_pos+1:].strip()

    # 적절한 분리점을 못 찾으면 첫 번째 콜론 사용
    pos = colon_positions[0]
    if 0 < pos <= max_name_length:
        return text[:pos].strip(), text[pos+1:].strip()

    return None, text


def _parse_roll20(soup, config):
    """Roll20 Chat Archive HTML 파싱"""
    entries = []
    parsing_config = config.get('parsing', {})
    normalize = parsing_config.get('normalize_punctuation', True)
    scene_patterns = config.get('chapter', {}).get('scene_patterns', [])

    # Roll20 메시지는 class="message" 로 시작하는 div에 있음
    messages = soup.find_all('div', class_=lambda x: x and 'message' in x)

    # 이전 화자 추적 (연속 메시지 처리)
    last_speaker = None

    for msg in messages:
        class_list = msg.get('class', [])
        class_str = ' '.join(class_list) if isinstance(class_list, list) else str(class_list)

        # 숨김 메시지 스킵
        if 'hidden-message' in class_str:
            continue

        # 이름 추출: <span class="by">이름:</span> 또는 <strong>이름</strong>
        by_span = msg.find('span', class_='by')
        name_tag = None
        name = None
        if by_span:
            name_tag = by_span
            name_text = by_span.get_text(strip=True)
            # 콜론 제거
            name = name_text.rstrip(':').strip()
            # :: 처리 (시스템 메시지)
            if name == ':' or name == '::':
                name = 'System'
            last_speaker = name
        else:
            # <strong> 태그에서 이름 추출 (대체 Roll20 형식)
            strong_tag = msg.find('strong')
            if strong_tag:
                name_tag = strong_tag
                name_text = strong_tag.get_text(strip=True)
                name = name_text.rstrip(':').strip()
                if name == ':' or name == '::':
                    name = 'System'
                last_speaker = name

        # 내용 추출: 이름 태그 이후의 텍스트
        content = ''
        if name_tag:
            # 이름 태그 다음의 모든 텍스트 노드 가져오기
            for sibling in name_tag.next_siblings:
                if hasattr(sibling, 'get_text'):
                    content += sibling.get_text(separator=' ', strip=True) + ' '
                elif isinstance(sibling, str):
                    content += sibling.strip() + ' '
            content = content.strip()
            # content가 비어있으면 부모의 content div에서 추출
            if not content and name_tag.parent:
                parent_text = name_tag.parent.get_text(strip=True)
                # 이름 부분 제거
                name_with_colon = name_tag.get_text(strip=True)
                if parent_text.startswith(name_with_colon):
                    content = parent_text[len(name_with_colon):].lstrip(':').strip()
        else:
            # by span이 없으면 전체 텍스트 (desc 메시지 등)
            # avatar, tstamp, spacer 등은 제외
            for child in msg.children:
                child_class = child.get('class', []) if hasattr(child, 'get') else []
                child_class_str = ' '.join(child_class) if isinstance(child_class, list) else str(child_class)
                if any(skip in child_class_str for skip in ['avatar', 'tstamp', 'spacer']):
                    continue
                if hasattr(child, 'get_text'):
                    content += child.get_text(separator=' ', strip=True) + ' '
                elif isinstance(child, str):
                    content += child.strip() + ' '
            content = content.strip()

        if not content:
            continue

        # 내용에서 선행 콜론 제거 (예: ": 대사내용")
        content = re.sub(r'^:\s*', '', content)

        # 노말라이제이션
        if normalize:
            content = normalize_punctuation(content)
        content = content.replace('＞', '→')

        # 엔트리 타입 결정
        entry_type = 'dialogue'

        # desc 클래스: 나레이션/설명
        if 'desc' in class_str:
            entry_type = 'narration'
            if not name:
                name = ''

        # 이름이 없는 연속 메시지: 이전 화자 사용
        if not name and 'general' in class_str and last_speaker:
            name = last_speaker

        # 장면 마커 감지
        if is_scene_marker(content, scene_patterns):
            entry_type = 'system'

        # 시스템 메시지
        if name == 'System' or name == '::':
            entry_type = 'system'
            name = ''

        # 다이스 롤 감지
        if is_dice_roll(content):
            entry_type = 'dice'

        # GM/KP/DM 등 나레이터 확인
        if name and is_narration_user(name, config):
            entry_type = 'narration'

        entry = {
            'type': entry_type,
            'name': name or '',
            'content': content,
            'raw': content,
            'image': None,
            'channel': None
        }
        entries.append(entry)

    return entries


def parse_log(html_content: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
    soup = BeautifulSoup(html_content, 'html.parser')
    entries = []

    parsing_config = config.get('parsing', {})
    skip_channels = parsing_config.get('skip_channels', [])
    name_max_length = parsing_config.get('name_max_length', 50)
    normalize = parsing_config.get('normalize_punctuation', True)
    scene_patterns = config.get('scene_detection', {}).get('patterns', [])

    # Roll20 형식 감지: id="textchat" 또는 class="message" 확인
    is_roll20 = soup.find(id='textchat') is not None or soup.find(class_='message') is not None

    # Roll20 형식 처리
    if is_roll20:
        entries = _parse_roll20(soup, config)
        return entries

    for element in soup.find_all(['p', 'div']):
        if element.find(['p', 'div']):
            continue

        # Cocofolia 형식: span class에서 이름 추출
        name_from_class = None
        content_from_spans = None
        for span in element.find_all('span'):
            class_attr = span.get('class', [])
            class_str = ' '.join(class_attr) if isinstance(class_attr, list) else str(class_attr)

            # firing_name_XXX 패턴에서 이름 추출
            name_match = re.search(r'firing_name_(\S+)', class_str)
            if name_match:
                name_from_class = name_match.group(1)
                # 같은 span에 firing_firing도 있으면 여기서 내용 추출 (일반적인 코코포리아 형식)
                if 'firing_firing' in class_str:
                    content_from_spans = span.get_text(strip=True)

            # 별도 span에 firing_firing만 있는 경우 (내용 전용 span)
            elif 'firing_firing' in class_str:
                content_from_spans = span.get_text(strip=True)

        # Cocofolia 형식이면 이름과 내용 분리
        if name_from_class:
            name = name_from_class.replace('_', ' ')  # 언더스코어를 공백으로
            if content_from_spans is not None:
                text = content_from_spans.lstrip(':').strip()
            else:
                # content가 없으면 전체 텍스트에서 추출
                text = element.get_text(separator=' ', strip=True).lstrip(':').strip()
        else:
            text = element.get_text(separator=' ', strip=True)
            name = None

        for channel in skip_channels:
            text = re.sub(re.escape(channel) + r'\s*', '', text)
        text = text.strip()

        # 채널 접두사 제거
        text, channel = strip_channel_prefix(text)

        # 채널이 skip_channels에 있으면 건너뛰기
        if channel and channel in skip_channels:
            continue

        if not text:
            continue

        text = text.replace('＞', '→')
        if normalize:
            text = normalize_punctuation(text)

        img_filename, text = extract_image_markers(text, config)
        entry = {'type': 'dialogue', 'name': name or '', 'content': text, 'raw': text, 'image': img_filename, 'channel': channel}

        if img_filename and not text.strip():
            entry['type'] = 'image'
            entry['content'] = img_filename
            entries.append(entry)
            continue

        if not text:
            continue

        # 장면 마커 감지
        if is_scene_marker(text, scene_patterns):
            entry['type'] = 'scene'
            entry['content'] = text
            if name:
                entry['name'] = name
            entries.append(entry)
            continue

        if text.lower().startswith('system :') or text.lower().startswith('system:'):
            content = re.sub(r'^system\s*:\s*', '', text, flags=re.IGNORECASE)
            entry['type'] = 'system'
            entry['content'] = content.strip()
        elif name:
            # Cocofolia 형식으로 이미 이름이 추출된 경우
            if is_narration_user(name, config):
                entry['type'] = 'narration'
                entry['name'] = name
            elif is_dice_roll(text):
                entry['type'] = 'dice'
                entry['name'] = name
            elif text.startswith('《') or text.count('|') >= 2:
                entry['type'] = 'effect'
                entry['name'] = name
            else:
                entry['type'] = 'dialogue'
                entry['name'] = name
        elif ':' in text:
            # 지능적 이름/내용 분리
            parsed_name, content = smart_split_name_content(text, name_max_length)

            if parsed_name and parsed_name.lower() != 'system':
                entry['name'] = parsed_name
                entry['content'] = content

                if is_narration_user(parsed_name, config):
                    entry['type'] = 'narration'
                else:
                    custom_type = match_custom_style(content, config)
                    if custom_type:
                        entry['type'] = custom_type
                    elif is_dice_roll(content):
                        entry['type'] = 'dice'
                    elif content.startswith('《') or content.count('|') >= 2:
                        entry['type'] = 'effect'
                    else:
                        entry['type'] = 'dialogue'

        entries.append(entry)

    return entries

# ============ 필터링 ============
def filter_entries(entries: List[Dict], config: Dict[str, Any]) -> List[Dict]:
    content_config = config.get('content', {})
    filtered = []
    
    for entry in entries:
        t = entry['type']
        if t == 'dice' and not content_config.get('include_dice', True):
            continue
        if t == 'system' and not content_config.get('include_system', True):
            patterns = config.get('chapter', {}).get('scene_patterns', [])
            if not is_scene_marker(entry.get('content', ''), patterns):
                continue
        if t == 'effect' and not content_config.get('include_effects', True):
            continue
        filtered.append(entry)
    
    return filtered

# ============ 대사 병합 ============
def merge_consecutive_dialogues(entries: List[Dict], config: Dict[str, Any]) -> List[Dict]:
    dialogue_config = config.get('dialogue', {})
    if not dialogue_config.get('merge_consecutive', False):
        return entries
    
    separator = dialogue_config.get('merge_separator', '\n')
    max_merge = dialogue_config.get('merge_max', 5)
    
    merged = []
    i = 0
    
    while i < len(entries):
        entry = entries[i].copy()
        
        if entry['type'] == 'dialogue' and entry.get('name'):
            same_speaker = [entry['content']]
            j = i + 1
            merge_count = 1
            
            while j < len(entries) and (max_merge == 0 or merge_count < max_merge):
                next_entry = entries[j]
                if (next_entry['type'] == 'dialogue' and 
                    next_entry.get('name', '').lower() == entry['name'].lower()):
                    same_speaker.append(next_entry['content'])
                    j += 1
                    merge_count += 1
                else:
                    break
            
            if len(same_speaker) > 1:
                entry['content'] = separator.join(same_speaker)
                i = j
            else:
                i += 1
        else:
            i += 1
        
        merged.append(entry)
    
    return merged

# ============ 장면 분할 ============
def split_into_scenes(entries: List[Dict], config: Dict[str, Any]) -> List[Dict]:
    chapter_config = config.get('chapter', {})
    split_mode = chapter_config.get('split_mode', 'scene')
    
    if split_mode == 'none':
        return [{'title': None, 'entries': entries}]
    
    if split_mode == 'count':
        return split_by_count(entries, chapter_config)
    
    return split_by_scene(entries, chapter_config)

def split_by_count(entries, chapter_config):
    per_chapter = chapter_config.get('entries_per_chapter', 300)
    title_format = chapter_config.get('title_format', '장면 {n}')
    scenes = []
    for i in range(0, len(entries), per_chapter):
        chunk = entries[i:i+per_chapter]
        scenes.append({'title': title_format.format(n=len(scenes)+1), 'entries': chunk})
    return scenes

def split_by_scene(entries, chapter_config):
    patterns = chapter_config.get('scene_patterns', [])
    min_entries = chapter_config.get('min_scene_entries', 10)
    title_format = chapter_config.get('title_format', '장면 {n}')
    extract_title = chapter_config.get('extract_scene_title', True)

    if not patterns:
        return [{'title': None, 'entries': entries}]

    scenes = []
    current_scene = {'title': None, 'entries': []}
    scene_count = 0

    for entry in entries:
        content = entry.get('content', '')
        entry_type = entry['type']

        # 장면 마커 체크: system 또는 narration(GM/KP/DM) 타입에서 확인
        is_scene = False
        if entry_type == 'system' and is_scene_marker(content, patterns):
            is_scene = True
        elif entry_type == 'narration' and is_scene_marker(content, patterns):
            is_scene = True

        if is_scene:
            if current_scene['entries']:
                if len(current_scene['entries']) >= min_entries or not scenes:
                    scenes.append(current_scene)
                elif scenes:
                    scenes[-1]['entries'].extend(current_scene['entries'])

            scene_count += 1
            scene_title = extract_scene_title(content) if extract_title else None
            if not scene_title:
                scene_title = title_format.format(n=scene_count)

            current_scene = {'title': scene_title, 'entries': []}
        else:
            current_scene['entries'].append(entry)

    if current_scene['entries']:
        if not current_scene['title']:
            current_scene['title'] = title_format.format(n=scene_count + 1) if scene_count else None
        scenes.append(current_scene)

    return scenes if scenes else [{'title': None, 'entries': entries}]

# ============ CSS ============
def generate_css(config, embedded_fonts=None):
    fonts = config.get('fonts', {})
    style = config.get('style', {})

    body_font = fonts.get('body_font', 'serif')
    name_font = fonts.get('name_font', 'sans-serif')

    # 비주얼 에디터 스타일 설정
    body_bg = style.get('body_bg', '#ffffff')
    body_text = style.get('body_text', '#1a1a1a')
    body_font_size = style.get('body_font_size', 14)
    name_color = style.get('name_color', '#2d2d2d')
    name_bold = 'bold' if style.get('name_bold', True) else 'normal'
    line_height = style.get('visual_line_height', 1.6)

    # 기타 스타일 설정
    dice_color = style.get('dice_color', '#888')
    system_color = style.get('system_color', '#666')
    effect_bg = style.get('effect_bg', '#f5f5f5')
    effect_border = style.get('effect_border', '#ccc')

    css_parts = [f"@import url('{fonts.get('pretendard_cdn', '')}');"]

    if embedded_fonts:
        for font_path in embedded_fonts.get('all', []):
            family_name = get_font_family_name(font_path)
            css_parts.append(f"@font-face {{ font-family: '{family_name}'; src: url('fonts/{font_path.name}'); }}")
        if embedded_fonts.get('body'):
            body_font = f"'{get_font_family_name(embedded_fonts['body'])}', {body_font}"
        if embedded_fonts.get('name'):
            name_font = f"'{get_font_family_name(embedded_fonts['name'])}', {name_font}"

    css_parts.append(f'''
body {{ font-family: {body_font}; font-size: {body_font_size}px; line-height: {line_height}; color: {body_text}; background-color: {body_bg}; margin: 1em; text-align: justify; }}
.dialogue {{ margin: 0.3em 0; line-height: 1.5; }}
.dialogue .name {{ font-family: {name_font}; font-weight: {name_bold}; color: {name_color}; font-size: 1em; margin-right: 0.8em; }}
.narration {{ font-size: 0.95em; line-height: 1.7; margin: 0.8em 0; padding-left: 1.5em; }}
.dice {{ font-family: {name_font}; font-size: 0.75em; color: {dice_color}; margin: 0.2em 0; padding-left: 1em; }}
.system {{ font-family: {name_font}; font-size: 0.85em; color: {system_color}; margin: 1em 0; padding-left: 0.5em; }}
.effect {{ font-family: {name_font}; font-size: 0.85em; color: #444; background: {effect_bg}; border-left: 3px solid {effect_border}; padding: 0.6em 1em; margin: 0.3em 0 0.3em 1em; }}
.whisper {{ font-size: 0.92em; font-style: italic; color: #666; padding-left: 1.5em; }}
.highlight {{ background: #fffde7; padding: 0.3em 0.5em; }}
.scene-title {{ font-family: {name_font}; font-weight: 700; font-size: 1em; margin: 1.5em 0 0.8em 0; }}
h1 {{ font-family: {name_font}; font-size: 1.4em; font-weight: 700; text-align: center; margin: 2em 0 1em 0; }}
.chapter-ornament {{ text-align: center; color: #bbb; margin: 0.5em 0 1.5em 0; font-size: 0.85em; }}
.cover {{ text-align: center; padding: 20% 10%; }}
.cover h1 {{ font-size: 2em; margin-bottom: 0.5em; }}
.toc {{ margin: 2em 0; }}
.toc h2 {{ text-align: center; margin-bottom: 1em; }}
.toc ul {{ list-style: none; padding: 0; }}
.toc li {{ margin: 0.5em 0; }}
.toc a {{ color: inherit; text-decoration: none; }}
.illustration {{ text-align: center; margin: 1.5em 0; }}
.illustration img {{ max-width: 100%; }}
.caption {{ font-size: 0.8em; color: #666; margin-top: 0.5em; }}
''')

    # 커스텀 CSS 추가
    custom_css = config.get('custom_css', '')
    if custom_css and custom_css.strip():
        css_parts.append('\n/* Custom CSS */\n' + custom_css)

    return '\n'.join(css_parts)

# ============ EPUB ============
def entries_to_html(entries, chapter_title=None, config=None):
    style = config.get('style', {}) if config else {}
    narration_prefix = style.get('narration_prefix', '＿')
    scene_marker = style.get('scene_marker', '■')
    images_config = config.get('images', {}) if config else {}
    
    html_parts = []
    
    if chapter_title:
        title_has_marker = chapter_title.startswith('■') or chapter_title.startswith('●')
        if scene_marker and not title_has_marker:
            html_parts.append(f'<p class="scene-title"><span class="scene-marker">{escape_html(scene_marker)}</span> {escape_html(chapter_title)}</p>')
        else:
            html_parts.append(f'<p class="scene-title">{escape_html(chapter_title)}</p>')
    
    for entry in entries:
        t = entry['type']
        content = entry.get('content', '')
        name = entry.get('name', '')
        img = entry.get('image')
        
        if img:
            img_path = find_image_file(img, config)
            if img_path:
                # WebP→JPEG 변환 시 파일명 반영
                convert_webp = images_config.get('convert_webp', True)
                ref_name = img_path.name
                if convert_webp and str(img_path).lower().endswith('.webp'):
                    ref_name = img_path.stem + '.jpg'
                html_parts.append(f'<div class="illustration"><img src="images/{ref_name}" alt="{escape_html(img)}"/>')
                if images_config.get('show_caption', True):
                    html_parts.append(f'<p class="caption">{escape_html(img)}</p>')
                html_parts.append('</div>')
        
        if t == 'image' or not content or not content.strip():
            continue
        
        content_escaped = escape_html(content).replace('\n', '<br/>')
        name_escaped = escape_html(name) if name else ''
        
        if t == 'dialogue':
            if name_escaped:
                html_parts.append(f'<p class="dialogue"><span class="name">{name_escaped}</span><span class="content">{content_escaped}</span></p>')
            else:
                html_parts.append(f'<p class="dialogue">{content_escaped}</p>')
        elif t == 'effect':
            if name_escaped:
                html_parts.append(f'<p class="dialogue"><span class="name">{name_escaped}</span></p>')
            html_parts.append(f'<div class="effect">{content_escaped}</div>')
        elif t == 'narration':
            prefix = narration_prefix if narration_prefix else ''
            html_parts.append(f'<p class="narration">{prefix}{content_escaped}</p>')
        elif t == 'dice':
            if name_escaped:
                html_parts.append(f'<p class="dice">{name_escaped} : {content_escaped}</p>')
            else:
                html_parts.append(f'<p class="dice">{content_escaped}</p>')
        elif t == 'system':
            html_parts.append(f'<p class="system">{content_escaped}</p>')
        elif t == 'whisper':
            html_parts.append(f'<p class="whisper">{content_escaped}</p>')
        elif t == 'highlight':
            html_parts.append(f'<p class="highlight">{content_escaped}</p>')
        else:
            html_parts.append(f'<p class="dialogue">{content_escaped}</p>')
    
    return '\n'.join(html_parts)

def create_cover_html(config, title, author):
    cover_config = config.get('cover', {})
    bg_color = cover_config.get('background_color', '#1a1a1a')
    title_color = cover_config.get('title_color', '#ffffff')
    subtitle = cover_config.get('subtitle', '')
    
    html = f'<div class="cover" style="background-color: {bg_color}; color: {title_color}; min-height: 100vh;">'
    if cover_config.get('title_on_cover', True):
        html += f'<h1 style="color: {title_color};">{escape_html(title)}</h1>'
    if subtitle:
        html += f'<p class="subtitle">{escape_html(subtitle)}</p>'
    if cover_config.get('author_on_cover', True):
        html += f'<p class="author">{escape_html(author)}</p>'
    html += '</div>'
    return html

def create_toc_html(scenes, config):
    toc_title = config.get('toc', {}).get('title', '목차')
    html = f'<div class="toc"><h2>{escape_html(toc_title)}</h2><ul>'
    for idx, scene in enumerate(scenes):
        title = scene.get('title') or f"장면 {idx + 1}"
        html += f'<li><a href="chapter_{idx+1}.xhtml">{escape_html(title)}</a></li>'
    html += '</ul></div>'
    return html

def create_epub(entries, output_path, config, title="TRPG 리플레이", author=None, progress_callback=None):
    if progress_callback:
        progress_callback(0, 100, "EPUB 생성 준비 중...")
    book = epub.EpubBook()
    
    if author is None:
        author = config.get('metadata', {}).get('author', 'GM')
    language = config.get('metadata', {}).get('language', 'ko')
    
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language(language)
    book.add_author(author)
    
    embedded_fonts = get_font_files(config)
    
    for font_path in embedded_fonts.get('all', []):
        mime_type = 'font/ttf' if font_path.suffix == '.ttf' else 'font/otf'
        with open(font_path, 'rb') as f:
            book.add_item(epub.EpubItem(uid=f"font_{font_path.stem}", file_name=f"fonts/{font_path.name}",
                                        media_type=mime_type, content=f.read()))
    
    css_content = generate_css(config, embedded_fonts)
    css = epub.EpubItem(uid="style", file_name="style/main.css", media_type="text/css", content=css_content)
    book.add_item(css)
    
    epub_chapters = []
    
    cover_config = config.get('cover', {})
    if cover_config.get('include', True):
        cover_image = cover_config.get('image', '')
        if cover_image:
            img_path = find_image_file(cover_image, config)
            if img_path:
                with open(img_path, 'rb') as f:
                    book.set_cover(img_path.name, f.read())
        
        cover_html = create_cover_html(config, title, author)
        cover_chapter = epub.EpubHtml(title='표지', file_name='cover.xhtml', lang=language)
        cover_chapter.set_content(f'''<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>표지</title><link rel="stylesheet" type="text/css" href="style/main.css"/></head>
<body>{cover_html}</body></html>''')
        cover_chapter.add_item(css)
        book.add_item(cover_chapter)
        epub_chapters.append(cover_chapter)
    
    if progress_callback:
        progress_callback(10, 100, "이미지 처리 중...")
    images_added = set()
    optimized_names = {}  # img_path -> optimized filename
    for entry in entries:
        img = entry.get('image')
        if img:
            img_path = find_image_file(img, config)
            if img_path and str(img_path) not in images_added:
                img_data, mime_type, opt_name = optimize_image(img_path, config)
                img_item = epub.EpubItem(uid=f"img_{img_path.stem}", file_name=f"images/{opt_name}",
                                        media_type=mime_type, content=img_data)
                book.add_item(img_item)
                images_added.add(str(img_path))
                optimized_names[str(img_path)] = opt_name
    
    if progress_callback:
        progress_callback(20, 100, "장면 분할 중...")
    scenes = split_into_scenes(entries, config)
    logger.info("Split into %d scenes", len(scenes))
    
    toc_config = config.get('toc', {})
    if toc_config.get('include', True) and len(scenes) > 1:
        toc_html = create_toc_html(scenes, config)
        toc_chapter = epub.EpubHtml(title=toc_config.get('title', '목차'), file_name='toc_page.xhtml', lang=language)
        toc_chapter.set_content(f'''<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>목차</title><link rel="stylesheet" type="text/css" href="style/main.css"/></head>
<body>{toc_html}</body></html>''')
        toc_chapter.add_item(css)
        book.add_item(toc_chapter)
        epub_chapters.append(toc_chapter)
    
    for idx, scene in enumerate(scenes):
        scene_title = scene.get('title')
        scene_entries = scene.get('entries', [])
        logger.debug("Scene '%s': %d entries", scene_title or f'장면 {idx+1}', len(scene_entries))
        if progress_callback:
            pct = 30 + int(60 * idx / max(len(scenes), 1))
            progress_callback(pct, 100, f"장면 {idx+1}/{len(scenes)} 생성 중...")
        
        content_html = entries_to_html(scene_entries, scene_title, config)
        
        chapter = epub.EpubHtml(title=scene_title or f"장면 {idx+1}", file_name=f'chapter_{idx+1}.xhtml', lang=language)
        chapter.set_content(f'''<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{escape_html(scene_title or title)}</title>
<link rel="stylesheet" type="text/css" href="style/main.css"/></head>
<body>{content_html}</body></html>''')
        chapter.add_item(css)
        book.add_item(chapter)
        epub_chapters.append(chapter)
    
    book.toc = epub_chapters
    book.spine = ['nav'] + epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    if progress_callback:
        progress_callback(90, 100, "EPUB 저장 중...")
    # 원자적 쓰기: 임시 파일에 작성 후 rename
    import tempfile
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.epub', dir=os.path.dirname(output_path) or '.')
    os.close(tmp_fd)
    try:
        epub.write_epub(tmp_path, book)
        if os.path.exists(output_path):
            os.replace(tmp_path, output_path)
        else:
            os.rename(tmp_path, output_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise
    if progress_callback:
        progress_callback(100, 100, "EPUB 생성 완료")
    logger.info("EPUB created: %s", output_path)
    return output_path

# ============ DOCX ============
def set_run_font(run, font_name, size_pt=None, bold=False, italic=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = hex_to_rgb(color) if isinstance(color, str) else color

def add_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing=1.0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def create_docx(entries, output_path, config, title="TRPG 리플레이", author=None, progress_callback=None):
    if progress_callback:
        progress_callback(0, 100, "DOCX 생성 준비 중...")
    doc = Document()
    
    if author is None:
        author = config.get('metadata', {}).get('author', 'GM')
    
    fonts = config.get('fonts', {})
    embedded = get_font_files(config)
    fallback = fonts.get('docx_fallback', {})
    style_config = config.get('style', {})
    narration_prefix = style_config.get('narration_prefix', '＿')
    
    body_font = get_font_family_name(embedded['body']) if embedded.get('body') else fallback.get('body', '맑은 고딕')
    name_font = get_font_family_name(embedded['name']) if embedded.get('name') else fallback.get('name', '맑은 고딕')
    
    cover_config = config.get('cover', {})
    if cover_config.get('include', True):
        cover_image = cover_config.get('image', '')
        if cover_image:
            img_path = find_image_file(cover_image, config)
            if img_path:
                img_data, _, _ = optimize_image(img_path, config)
                doc.add_picture(io.BytesIO(img_data), width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if cover_config.get('title_on_cover', True):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            set_run_font(title_run, name_font, size_pt=28, bold=True)
            add_paragraph_spacing(title_para, before_pt=72, after_pt=24)
        
        if cover_config.get('author_on_cover', True):
            auth_para = doc.add_paragraph()
            auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_run = auth_para.add_run(author)
            set_run_font(auth_run, name_font, size_pt=12, color='#888888')
        
        doc.add_page_break()
    
    scenes = split_into_scenes(entries, config)

    for scene_idx, scene in enumerate(scenes):
        scene_title = scene.get('title')
        scene_entries = scene.get('entries', [])
        if progress_callback:
            pct = 10 + int(80 * scene_idx / max(len(scenes), 1))
            progress_callback(pct, 100, f"장면 {scene_idx+1}/{len(scenes)} 생성 중...")
        
        if scene_title:
            if scene_idx > 0:
                doc.add_page_break()
            
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            title_has_marker = scene_title.startswith('■')
            display_title = scene_title if title_has_marker else f"■ {scene_title}"
            
            title_run = title_para.add_run(display_title)
            set_run_font(title_run, name_font, size_pt=14, bold=True)
            add_paragraph_spacing(title_para, before_pt=24, after_pt=16)
        
        for entry in scene_entries:
            t = entry['type']
            content = entry.get('content', '')
            name = entry.get('name', '')
            img = entry.get('image')
            
            if img:
                img_path = find_image_file(img, config)
                if img_path:
                    img_data, _, _ = optimize_image(img_path, config)
                    doc.add_picture(io.BytesIO(img_data), width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if t == 'image' or not content.strip():
                continue
            
            para = doc.add_paragraph()
            
            if t == 'dialogue':
                if name:
                    name_run = para.add_run(f"{name}   ")
                    set_run_font(name_run, name_font, size_pt=11, bold=True)
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=11)
                add_paragraph_spacing(para, before_pt=2, after_pt=2, line_spacing=1.3)
            
            elif t == 'narration':
                prefix = narration_prefix if narration_prefix else ''
                content_run = para.add_run(f"{prefix}{content}")
                set_run_font(content_run, body_font, size_pt=10.5)
                para.paragraph_format.left_indent = Inches(0.3)
                add_paragraph_spacing(para, before_pt=8, after_pt=8, line_spacing=1.4)
            
            elif t == 'dice':
                dice_text = f"{name} : {content}" if name else content
                content_run = para.add_run(dice_text)
                set_run_font(content_run, name_font, size_pt=9, color='#888888')
                para.paragraph_format.left_indent = Inches(0.2)
                add_paragraph_spacing(para, before_pt=1, after_pt=1, line_spacing=1.2)
            
            elif t == 'system':
                content_run = para.add_run(content)
                set_run_font(content_run, name_font, size_pt=10, color='#666666')
                add_paragraph_spacing(para, before_pt=12, after_pt=12)
            
            elif t == 'effect':
                if name:
                    name_run = para.add_run(f"{name}\n")
                    set_run_font(name_run, name_font, size_pt=10, bold=True)
                content_run = para.add_run(content)
                set_run_font(content_run, name_font, size_pt=9, color='#444444')
                para.paragraph_format.left_indent = Inches(0.3)
                add_paragraph_spacing(para, before_pt=4, after_pt=4, line_spacing=1.2)
            
            elif t == 'whisper':
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=10, italic=True, color='#666666')
                para.paragraph_format.left_indent = Inches(0.4)
            
            else:
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=11)
    
    if progress_callback:
        progress_callback(95, 100, "DOCX 저장 중...")
    # 원자적 쓰기: 임시 파일에 작성 후 rename
    import tempfile
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=os.path.dirname(output_path) or '.')
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        if os.path.exists(output_path):
            os.replace(tmp_path, output_path)
        else:
            os.rename(tmp_path, output_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise
    if progress_callback:
        progress_callback(100, 100, "DOCX 생성 완료")
    logger.info("DOCX created: %s", output_path)
    return output_path

# ============ ConversionEngine 클래스 ============
class ConversionEngine:
    """변환 엔진 클래스 - GUI에서 사용"""

    def __init__(self, config=None):
        self.config = config or load_config()

    def parse_file(self, file_path):
        """파일 파싱"""
        from core.text_parser import parse_file as text_parse_file
        return text_parse_file(file_path, self.config)

    def parse_html(self, html_content):
        """HTML 콘텐츠 파싱"""
        return parse_log(html_content, self.config)

    def filter(self, entries):
        """엔트리 필터링"""
        return filter_entries(entries, self.config)

    def merge(self, entries):
        """연속 대사 병합"""
        return merge_consecutive_dialogues(entries, self.config)

    def split_scenes(self, entries):
        """장면 분할"""
        return split_into_scenes(entries, self.config)

    def create_epub(self, entries, output_path, title="TRPG 리플레이", author=None, progress_callback=None):
        """EPUB 생성"""
        return create_epub(entries, output_path, self.config, title, author, progress_callback=progress_callback)

    def create_docx(self, entries, output_path, title="TRPG 리플레이", author=None, progress_callback=None):
        """DOCX 생성"""
        return create_docx(entries, output_path, self.config, title, author, progress_callback=progress_callback)

    def create_pdf(self, entries, output_path, title="TRPG 리플레이", author=None):
        """PDF 생성"""
        try:
            from core.pdf_generator import create_pdf as pdf_create, PDF_AVAILABLE
            if PDF_AVAILABLE:
                return pdf_create(entries, output_path, self.config, title, author)
            return None
        except ImportError:
            return None

    def convert_file(self, input_path, output_path=None, title=None, author=None, format=None, progress_callback=None):
        """단일 파일 변환"""
        return convert(input_path, output_path, title, author, self.config, format, progress_callback=progress_callback)

    def batch_convert(self, input_files, output_dir=None, title_prefix=None, author=None, format=None, progress_callback=None):
        """배치 변환 - 여러 파일 한번에"""
        if output_dir:
            self.config['paths']['output_dir'] = output_dir

        results = []
        for i, file_path in enumerate(input_files, 1):
            try:
                base = os.path.splitext(os.path.basename(file_path))[0]
                title = f"{title_prefix} - {base}" if title_prefix else base
                if progress_callback:
                    progress_callback(i - 1, len(input_files), f"변환 중: {base}")
                result = convert(file_path, None, title, author, self.config, format, progress_callback=progress_callback)
                results.append({'file': file_path, 'success': True, 'outputs': result})
                logger.info("[%d/%d] OK: %s", i, len(input_files), base)
            except Exception as e:
                results.append({'file': file_path, 'success': False, 'error': str(e)})
                logger.error("[%d/%d] FAIL: %s: %s", i, len(input_files), base, e)

        return results


# ============ 메인 ============
def convert(input_path, output_path=None, title=None, author=None, config=None, format=None, progress_callback=None):
    if config is None:
        config = load_config()
    
    output_dir = config.get('paths', {}).get('output_dir', './export')
    os.makedirs(output_dir, exist_ok=True)
    
    base = os.path.splitext(os.path.basename(input_path))[0]
    if title is None:
        title = base
    
    if progress_callback:
        progress_callback(0, 100, "파일 읽기 중...")
    html_content = None
    for enc in ('utf-8', 'euc-kr', 'cp949', 'latin-1'):
        try:
            with open(input_path, 'r', encoding=enc) as f:
                html_content = f.read()
            break
        except UnicodeDecodeError:
            continue
    if html_content is None:
        raise UnicodeDecodeError('utf-8', b'', 0, 1, f"파일을 읽을 수 없습니다: {input_path}")

    logger.info("Input: %s", input_path)

    if progress_callback:
        progress_callback(10, 100, "파일 파싱 중...")
    entries = parse_log(html_content, config)
    logger.info("Parsed %d entries", len(entries))

    entries = filter_entries(entries, config)
    entries = merge_consecutive_dialogues(entries, config)

    stats = {}
    for e in entries:
        stats[e['type']] = stats.get(e['type'], 0) + 1
    logger.debug("Entry stats: %s", ' | '.join(f'{k} {v}' for k, v in sorted(stats.items())))
    
    output_format = format or config.get('output_format', 'both')
    results = []

    # EPUB 생성
    if output_format in ['epub', 'both', 'all']:
        if progress_callback:
            progress_callback(30, 100, "EPUB 생성 중...")
        epub_path = output_path or os.path.join(output_dir, f"{base}.epub")
        create_epub(entries, epub_path, config, title, author, progress_callback=progress_callback)
        results.append(epub_path)

    # DOCX 생성
    if output_format in ['docx', 'both', 'all']:
        if progress_callback:
            progress_callback(60, 100, "DOCX 생성 중...")
        docx_path = os.path.join(output_dir, f"{base}.docx")
        create_docx(entries, docx_path, config, title, author, progress_callback=progress_callback)
        results.append(docx_path)

    # PDF 생성
    if output_format in ['pdf', 'all']:
        try:
            from core.pdf_generator import create_pdf, PDF_AVAILABLE
            if PDF_AVAILABLE:
                pdf_path = os.path.join(output_dir, f"{base}.pdf")
                result = create_pdf(entries, pdf_path, config, title, author)
                if result:
                    results.append(pdf_path)
            else:
                logger.warning("PDF unavailable (reportlab not installed)")
        except ImportError:
            logger.warning("PDF module load failed")

    return results

def batch_convert(input_dir=None, output_dir=None, config=None, format=None, progress_callback=None):
    if config is None:
        config = load_config()
    
    if input_dir is None:
        input_dir = config.get('paths', {}).get('input_dir', './input')
    
    html_files = list(Path(input_dir).glob('*.html')) + list(Path(input_dir).glob('*.htm'))
    
    if not html_files:
        logger.warning("No HTML files in %s", input_dir)
        return []

    logger.info("Batch converting %d files", len(html_files))

    results = []
    for html_file in sorted(html_files):
        try:
            output_paths = convert(str(html_file), config=config, format=format)
            results.extend(output_paths)
        except Exception as e:
            logger.error("Conversion error: %s - %s", html_file, e)

    logger.info("Batch complete: %d files created", len(results))
    
    return results

if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='TRPG 로그 → EPUB/DOCX')
    parser.add_argument('input', nargs='?', help='입력 HTML')
    parser.add_argument('-o', '--output', help='출력 경로')
    parser.add_argument('-t', '--title', help='제목')
    parser.add_argument('-a', '--author', help='저자')
    parser.add_argument('-c', '--config', help='설정 파일')
    parser.add_argument('-f', '--format', choices=['epub', 'docx', 'both'])
    parser.add_argument('--batch', action='store_true')
    
    args = parser.parse_args()
    
    config = load_config(args.config)
    
    if args.batch or args.input is None:
        batch_convert(args.input, args.output, config, args.format)
    else:
        convert(args.input, args.output, args.title, args.author, config, args.format)
