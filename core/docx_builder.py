#!/usr/bin/env python3
"""
DOCX 빌더 — TRPG 로그 entries를 .docx 파일로 변환한다.

이 모듈은 core.engine의 파싱/유틸 헬퍼(get_font_files, find_image_file,
split_into_scenes 등)에 의존하지만, engine.py는 모듈 최하단에서
이 모듈을 re-export하므로 순환 import는 발생하지 않는다.
"""

from __future__ import annotations

import io
import logging
import os
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from core.parsers.fonts import get_font_family_name, get_font_files
from core.parsers.helpers import hex_to_rgb
from core.parsers.images import find_image_file, optimize_image
from core.parsers.pipeline import split_into_scenes

logger = logging.getLogger(__name__)


def set_run_font(run, font_name, size_pt=None, bold=False, italic=False, color=None):
    # Word 가 한글에 지정 폰트를 확실히 적용하도록 rFonts 의 네 속성 + hint 를 모두 지정.
    # w:eastAsia 만 채우면 Word 는 w:hint 가 없을 때 한글을 Latin(w:ascii) 폰트로 그리는
    # 경우가 있어 한글 폰트가 무시되고 기본 글꼴(맑은 고딕 등)로 fallback 됨.
    run.font.name = font_name
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:hint"), "eastAsia")
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = hex_to_rgb(color) if isinstance(color, str) else color


def add_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing=1.0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _coerce_float(value, fallback):
    """config 값이 문자열/None 일 수 있으므로 안전하게 float 로 변환한다."""
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(fallback)


def _set_para_shading(paragraph, fill_hex):
    """문단 배경색(w:shd)을 지정. fill_hex 가 falsy 면 아무것도 하지 않는다."""
    if not fill_hex:
        return
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill_hex).lstrip("#"))
    pPr.append(shd)


def _set_para_left_border(paragraph, color_hex, sz=18):
    """문단 왼쪽 테두리(w:pBdr > w:left). color_hex 가 falsy 면 아무것도 하지 않는다."""
    if not color_hex:
        return
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), str(color_hex).lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)


def _add_docx_divider(doc, config, name_font):
    """장면 구분선(장식)을 챕터 제목 아래에 추가. type 별로 텍스트/이미지/선."""
    d = (config or {}).get("divider", {}) or {}
    dtype = d.get("type", "텍스트/기호")
    if dtype == "사용 안 함":
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if dtype == "이미지":
        img_b64 = d.get("image", "")
        if not img_b64:
            return
        try:
            import base64

            data = base64.b64decode(img_b64)
            h_px = int(d.get("img_height", 40))
            run = para.add_run()
            # px → inch (96dpi). SVG 등 docx 가 못 읽는 형식이면 예외 → 건너뜀.
            run.add_picture(io.BytesIO(data), height=Inches(max(8, h_px) / 96.0))
        except Exception:
            return
    elif dtype == "선 스타일":
        width = int(d.get("line_width", 60))
        line_style = d.get("line_style", "실선")
        glyph = {"실선": "─", "점선": "┈", "이중선": "═"}.get(line_style, "─")
        try:
            thickness = int(d.get("line_thickness", 1))
        except (TypeError, ValueError):
            thickness = 1
        size_pt = 10 + min(max(thickness, 0), 4)
        run = para.add_run(glyph * max(3, width // 3))
        set_run_font(run, name_font, size_pt=size_pt, color=d.get("line_color", "#cccccc"))
    else:  # 텍스트/기호
        text = d.get("text", "")
        if not text.strip():
            return
        run = para.add_run(text)
        set_run_font(run, name_font, size_pt=11, color=d.get("color", "#888888"))
    add_paragraph_spacing(para, before_pt=4, after_pt=14)


def create_docx(
    entries, output_path, config, title="TRPG 리플레이", author=None, progress_callback=None
):
    if progress_callback:
        progress_callback(0, 100, "DOCX 생성 준비 중...")
    doc = Document()

    # 페이지 판형 / 여백 적용 — PDF 와 동일한 core.layout 헬퍼를 사용해
    # DOCX 와 PDF 의 레이아웃이 일치하도록 보장한다.
    try:
        from core.layout import get_page_format, get_page_margins_inch, parse_page_format

        page_w_mm, page_h_mm = parse_page_format(get_page_format(config))
        margins = get_page_margins_inch(config)
        section = doc.sections[0]
        section.page_width = Mm(page_w_mm)
        section.page_height = Mm(page_h_mm)
        section.top_margin = Inches(margins["top"])
        section.bottom_margin = Inches(margins["bottom"])
        section.left_margin = Inches(margins["left"])
        section.right_margin = Inches(margins["right"])
    except Exception as e:
        logger.warning("DOCX 페이지 설정 실패, 기본값 사용: %s", e)

    if author is None:
        author = config.get("metadata", {}).get("author", "GM")

    fonts = config.get("fonts", {})
    embedded = get_font_files(config)
    fallback = fonts.get("docx_fallback", {})
    style_config = config.get("style", {})
    narration_prefix = style_config.get("narration_prefix", "＿")

    body_font = (
        get_font_family_name(embedded["body"])
        if embedded.get("body")
        else fallback.get("body", "맑은 고딕")
    )
    name_font = (
        get_font_family_name(embedded["name"])
        if embedded.get("name")
        else fallback.get("name", "맑은 고딕")
    )

    # 사용자 스타일 설정을 EPUB 과 동일하게 DOCX 에도 반영한다.
    name_color = style_config.get("name_color", "#2d2d2d")
    name_bold = bool(style_config.get("name_bold", True))
    body_text = style_config.get("body_text", "#1a1a1a")
    dice_color = style_config.get("dice_color", "#888888")
    system_color = style_config.get("system_color", "#666666")
    effect_bg = style_config.get("effect_bg", "#f5f5f5")
    effect_border = style_config.get("effect_border", "#cccccc")
    dialogue_line_height = _coerce_float(style_config.get("dialogue_line_height", 1.5), 1.5)
    narration_line_height = _coerce_float(style_config.get("narration_line_height", 1.7), 1.7)
    narration_indent_em = _coerce_float(style_config.get("narration_indent", 1.5), 1.5)
    header_cfg = config.get("header", {}) or {}
    images_cfg = config.get("images", {}) or {}
    show_caption = bool(images_cfg.get("show_caption", True))

    cover_config = config.get("cover", {})
    if cover_config.get("include", True):
        cover_image = cover_config.get("image", "")
        if cover_image:
            img_path = find_image_file(cover_image, config)
            if img_path:
                img_data, _, _ = optimize_image(img_path, config)
                doc.add_picture(io.BytesIO(img_data), width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if cover_config.get("title_on_cover", True):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            set_run_font(title_run, name_font, size_pt=28, bold=True)
            add_paragraph_spacing(title_para, before_pt=72, after_pt=24)

        if cover_config.get("author_on_cover", True):
            auth_para = doc.add_paragraph()
            auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_run = auth_para.add_run(author)
            set_run_font(auth_run, name_font, size_pt=12, color="#888888")

        doc.add_page_break()

    scenes = split_into_scenes(entries, config)
    drop_cap = bool(config.get("style", {}).get("first_letter", False))

    for scene_idx, scene in enumerate(scenes):
        scene_title = scene.get("title")
        scene_entries = scene.get("entries", [])
        if progress_callback:
            pct = 10 + int(80 * scene_idx / max(len(scenes), 1))
            progress_callback(pct, 100, f"장면 {scene_idx + 1}/{len(scenes)} 생성 중...")

        if scene_title:
            # 'page_break' 설정이 켜져 있을 때만 챕터마다 페이지를 나눈다.
            # (이전엔 무조건 분할되어 UI 체크 해제가 무시됐다.)
            page_break_enabled = config.get("layout", {}).get("page_break", True)
            if scene_idx > 0 and page_break_enabled:
                doc.add_page_break()

            # 챕터 느낌의 큰 제목 — 사용자 설정(config['header'])을 따름
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ■ ▶ 같은 시각적 군더더기 마커는 제거. prefix/suffix 가 설정돼 있으면 추가
            base_title = scene_title.lstrip("■▶ ").strip() or scene_title
            display_title = (
                f"{header_cfg.get('prefix', '')}{base_title}{header_cfg.get('suffix', '')}"
            )

            title_run = title_para.add_run(display_title)
            try:
                header_size = float(header_cfg.get("size", 24))
            except (TypeError, ValueError):
                header_size = 24.0
            set_run_font(
                title_run,
                name_font,
                size_pt=header_size,
                bold=bool(header_cfg.get("bold", True)),
                color=header_cfg.get("color") or None,
            )
            if header_cfg.get("underline"):
                title_run.font.underline = True
            if header_cfg.get("box"):
                _set_para_shading(title_para, header_cfg.get("box_color", "#f5f5f5"))
            add_paragraph_spacing(title_para, before_pt=72, after_pt=36)

            # 장면 구분선(장식) — 제목 아래
            _add_docx_divider(doc, config, name_font)

        # 드롭캡: 이 챕터의 첫 나레이션 문단에만 적용.
        dropcap_pending = drop_cap

        for entry in scene_entries:
            t = entry["type"]
            content = entry.get("content", "")
            name = entry.get("name", "")
            img = entry.get("image")

            if img:
                img_path = find_image_file(img, config)
                if img_path:
                    img_data, _, _ = optimize_image(img_path, config)
                    doc.add_picture(io.BytesIO(img_data), width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 이미지 캡션 — EPUB 은 entry["image"] 값을 캡션으로 쓴다.
                    # 여기서는 명시적 caption 필드를 우선하고, image 타입이면 content,
                    # 그래도 없으면 EPUB 과 동일하게 image 값(img) 으로 폴백한다.
                    if show_caption:
                        caption = entry.get("caption")
                        if not caption and t == "image":
                            caption = entry.get("content")
                        if not caption:
                            caption = img
                        if caption and str(caption).strip():
                            cap_para = doc.add_paragraph()
                            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = cap_para.add_run(str(caption))
                            set_run_font(cap_run, body_font, size_pt=9, color="#666666")
                            add_paragraph_spacing(cap_para, before_pt=2, after_pt=8)

            if t == "image" or not (content and content.strip()):
                continue

            para = doc.add_paragraph()

            if t == "dialogue":
                if name:
                    name_run = para.add_run(f"{name}   ")
                    set_run_font(name_run, name_font, size_pt=11, bold=name_bold, color=name_color)
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=11, color=body_text)
                add_paragraph_spacing(
                    para, before_pt=2, after_pt=2, line_spacing=dialogue_line_height
                )

            elif t == "narration":
                prefix = narration_prefix if narration_prefix else ""
                if dropcap_pending:
                    from core.utils import split_first_grapheme

                    first, rest = split_first_grapheme(content.lstrip())
                    cap_run = para.add_run(first)
                    set_run_font(cap_run, name_font, size_pt=30, bold=True)
                    rest_run = para.add_run(rest)
                    set_run_font(rest_run, body_font, size_pt=10.5, color=body_text)
                    dropcap_pending = False
                else:
                    content_run = para.add_run(f"{prefix}{content}")
                    set_run_font(content_run, body_font, size_pt=10.5, color=body_text)
                para.paragraph_format.left_indent = Inches(narration_indent_em * 0.2)
                add_paragraph_spacing(
                    para, before_pt=8, after_pt=8, line_spacing=narration_line_height
                )

            elif t == "dice":
                dice_text = f"{name} : {content}" if name else content
                content_run = para.add_run(dice_text)
                set_run_font(content_run, name_font, size_pt=9, color=dice_color)
                para.paragraph_format.left_indent = Inches(0.2)
                add_paragraph_spacing(para, before_pt=1, after_pt=1, line_spacing=1.2)

            elif t == "system":
                content_run = para.add_run(content)
                set_run_font(content_run, name_font, size_pt=10, color=system_color)
                add_paragraph_spacing(para, before_pt=12, after_pt=12)

            elif t == "effect":
                if name:
                    name_run = para.add_run(f"{name}\n")
                    set_run_font(name_run, name_font, size_pt=10, bold=True)
                content_run = para.add_run(content)
                set_run_font(content_run, name_font, size_pt=9, color="#444444")
                para.paragraph_format.left_indent = Inches(0.3)
                _set_para_shading(para, effect_bg)
                _set_para_left_border(para, effect_border)
                add_paragraph_spacing(para, before_pt=4, after_pt=4, line_spacing=1.2)

            elif t == "scene_end":
                # 직전 챕터 마지막 줄에 작은 폰트로 가운데 정렬
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                content_run = para.add_run(content)
                set_run_font(content_run, name_font, size_pt=9, color="#999999")
                add_paragraph_spacing(para, before_pt=18, after_pt=6, line_spacing=1.0)

            elif t == "whisper":
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=10, italic=True, color="#666666")
                para.paragraph_format.left_indent = Inches(0.4)

            else:
                content_run = para.add_run(content)
                set_run_font(content_run, body_font, size_pt=11)

    if progress_callback:
        progress_callback(95, 100, "DOCX 저장 중...")

    # 원자적 쓰기: 임시 파일에 작성 후 rename
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(output_path) or ".")
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        if os.path.exists(output_path):
            os.replace(tmp_path, output_path)
        else:
            os.rename(tmp_path, output_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    if progress_callback:
        progress_callback(100, 100, "DOCX 생성 완료")
    logger.info("DOCX created: %s", output_path)
    return output_path
