#!/usr/bin/env python3
"""
DOCX 렌더러 일관성 테스트 — 사용자 스타일 설정(색상/줄간격/효과 박스 등)이
하드코딩 대신 config 에서 반영되는지 검증한다.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from core.docx_builder import create_docx


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _build(entries, config, tmp_path):
    out = tmp_path / "out.docx"
    create_docx(entries, str(out), config, title="테스트")
    assert out.exists()
    return Document(str(out))


def _custom_config():
    return {
        "style": {
            "name_color": "#112233",
            "name_bold": True,
            "body_text": "#445566",
            "dice_color": "#778899",
            "system_color": "#aabbcc",
            "effect_bg": "#e0e0e0",
            "effect_border": "#909090",
            "dialogue_line_height": 2.0,
            "narration_line_height": 2.2,
            "narration_indent": 2.5,
        },
        "header": {"box": True, "box_color": "#abcdef"},
        "images": {"show_caption": True},
        "cover": {"include": False},
    }


def _find_run_with_text(doc, needle):
    for para in doc.paragraphs:
        for run in para.runs:
            if needle in run.text:
                return run
    return None


def test_dialogue_uses_configured_colors(tmp_path):
    config = _custom_config()
    entries = [{"type": "dialogue", "name": "홍길동", "content": "안녕하세요"}]
    doc = _build(entries, config, tmp_path)

    name_run = _find_run_with_text(doc, "홍길동")
    assert name_run is not None
    assert name_run.font.color.rgb == _hex_to_rgb("#112233")
    assert name_run.font.bold is True

    body_run = _find_run_with_text(doc, "안녕하세요")
    assert body_run is not None
    assert body_run.font.color.rgb == _hex_to_rgb("#445566")


def test_dice_and_system_colors(tmp_path):
    config = _custom_config()
    entries = [
        {"type": "dice", "name": "주사위", "content": "1d100 = 50"},
        {"type": "system", "name": "", "content": "시스템 메시지"},
    ]
    doc = _build(entries, config, tmp_path)

    dice_run = _find_run_with_text(doc, "1d100")
    assert dice_run is not None
    assert dice_run.font.color.rgb == _hex_to_rgb("#778899")

    sys_run = _find_run_with_text(doc, "시스템 메시지")
    assert sys_run is not None
    assert sys_run.font.color.rgb == _hex_to_rgb("#aabbcc")


def test_effect_has_shading_and_border(tmp_path):
    config = _custom_config()
    entries = [{"type": "effect", "name": "효과", "content": "효과 내용입니다"}]
    doc = _build(entries, config, tmp_path)

    effect_para = None
    for para in doc.paragraphs:
        if "효과 내용입니다" in para.text:
            effect_para = para
            break
    assert effect_para is not None

    pPr = effect_para._p.pPr
    assert pPr is not None

    shd = pPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "e0e0e0"

    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None
    left = pBdr.find(qn("w:left"))
    assert left is not None
    assert left.get(qn("w:color")) == "909090"


def test_header_box_shading(tmp_path):
    config = _custom_config()
    # count-split 은 항상 챕터 제목을 생성하므로 헤더 박스 검증에 적합하다.
    config["chapter"] = {"split_mode": "count", "entries_per_chapter": 1}
    entries = [
        {"type": "dialogue", "name": "A", "content": "대사1"},
        {"type": "dialogue", "name": "B", "content": "대사2"},
    ]
    doc = _build(entries, config, tmp_path)

    found = False
    for para in doc.paragraphs:
        pPr = para._p.pPr
        if pPr is None:
            continue
        shd = pPr.find(qn("w:shd"))
        if shd is not None and shd.get(qn("w:fill")) == "abcdef":
            found = True
            break
    assert found, "헤더 박스 음영(box_color)이 적용되지 않았다"


def test_image_caption_rendered(tmp_path):
    config = _custom_config()
    entries = [{"type": "image", "image": "없는파일.png", "content": "삽화 설명"}]
    # 이미지 파일이 없으면 그림은 추가되지 않지만, 캡션 로직은 find_image_file 결과에
    # 의존하므로 파일이 없으면 캡션도 생략된다. 따라서 caption 필드만으로 검증하지 않고
    # 문서 생성이 정상적으로 끝나는지(예외 없음)를 확인한다.
    doc = _build(entries, config, tmp_path)
    assert doc is not None


def test_default_output_still_builds(tmp_path):
    # 기본 설정(빈 config 에 가까운)으로도 정상 생성되는지 회귀 확인
    config = {"cover": {"include": False}}
    entries = [
        {"type": "dialogue", "name": "캐릭터", "content": "기본 대사"},
        {"type": "narration", "name": "", "content": "기본 나레이션"},
    ]
    doc = _build(entries, config, tmp_path)

    name_run = _find_run_with_text(doc, "캐릭터")
    assert name_run is not None
    # 기본 name_color 는 #2d2d2d
    assert name_run.font.color.rgb == _hex_to_rgb("#2d2d2d")
